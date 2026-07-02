#!/usr/bin/env python3
"""Generate docs/WRITTEN_REPORT.docx from the markdown report."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "docs" / "WRITTEN_REPORT.docx"

TITLE = "Policy Intelligence and the Industrialization of Healthcare Payment Integrity"

SECTIONS = [
    (
        "Topic Concept Definition",
        """Healthcare payment integrity rests on a deceptively simple premise: claims must comply with a vast, continuously evolving body of written policy. That corpus includes Centers for Medicare & Medicaid Services (CMS) bulletins and transmittals, National Correct Coding Initiative (NCCI) manuals, Medicare Administrative Contractor (MAC) local coverage determinations, payer medical policies, code-set revisions, and increasingly, value-based payment (VBP) contract language. Historically, converting this prose into executable claim edits—software rules that fire at submission or adjudication—has been a manual, expert-intensive craft. Cotiviti built its market position on that craft: professional coders and payment-integrity specialists read policy, interpret intent, and encode logic into edit libraries that recover or prevent improper dollars.

Policy intelligence names the next industrialization of that competency. Rather than treating each policy revision as a bespoke consulting engagement, an agentic pipeline ingests source documents, summarizes material provisions, detects version-to-version changes, drafts structured payment-integrity edits with citations, and routes every draft to a human subject-matter expert (SME) before simulation or deployment. The output is not a narrative summary alone but machine-readable rule logic—JSON or YAML structures with pseudocode and SQL previews—that can hand off to enterprise rules engines such as Edifecs. The hackathon proof of concept (Policy Intelligence Pipeline; Cotiviti, 2026) implements this as a four-stage workflow: ingest and summarize, change detection, rule drafting, and SME sign-off with a full audit trail.

This concept differs from generic clinical documentation improvement or prior-authorization automation. Payment-integrity policy intelligence targets financial compliance: modifier requirements, procedure-to-procedure bundling, medically unlikely edit thresholds, coverage frequency limits, and contract settlement thresholds. It augments—not replaces—professional judgment. Agents accelerate reading, diffing, and first-draft authoring; SMEs retain authority over approval, consistent with emerging industry guidance on agentic AI with human oversight (Edifecs, 2025).""",
    ),
    (
        "Relevant Trends Analysis",
        """Several converging forces make policy intelligence strategically urgent for payment-integrity vendors.

Agentic AI and large language models. Generative and agent-based systems can process unstructured policy PDFs, extract codes and effective dates, and propose structured edits far faster than manual review alone. The risk—hallucinated codes, misread effective dates, or over-broad rules—makes human-in-the-loop design non-negotiable rather than optional. Regulators and payers will expect demonstrable audit trails linking every deployed edit to source text and approver identity.

Prepay versus postpay shift. Post-payment recovery generates revenue but irritates providers, consumes appeal capacity, and captures only a fraction of leakage. Prepay edits applied at claim intake or real-time adjudication stop improper payment before funds leave the plan. CMS itself operationalizes prepay logic through NCCI procedure-to-procedure and medically unlikely edits updated quarterly—and occasionally through mid-quarter replacement files with retroactive effective dates (Centers for Medicare & Medicaid Services [CMS], n.d.). Payers seeking lower abrasion and faster yield favor partners who industrialize prepay rule deployment.

CMS policy velocity. Revenue integrity teams must track quarterly NCCI file releases, annual policy manual revisions, and ad hoc bulletins. NAHRI (2025) notes that practitioners must validate dates of service against the edit version in force—a operational burden that scales poorly without automated change detection. A pipeline that ingests v2024 and v2025 NCCI policy text, ranks diffs by dollar impact, and drafts executable edits directly addresses this velocity gap.

Human-in-the-loop governance. Section 3 assessment frameworks for health care content management explicitly expect conversion of written policy into rules with professional oversight. Fully autonomous rule deployment is neither compliant nor credible in Medicare-facing workflows. The correct model pairs agent speed with SME gates on simulation and export—the pattern encoded in the proof-of-concept pipeline.

VBP contract growth. Shared savings, quality thresholds, and downside risk arrangements embed payment logic in contract prose rather than in fee schedules alone. Extending the same ingest-diff-draft-review pattern to payer–provider contracts opens a Phase 2 product surface: contract administration and settlement logic industrialization, adjacent to but distinct from traditional claim edits.""",
    ),
    (
        "Opportunities and Threats for Cotiviti",
        """Opportunities. Cotiviti's historical strength—translating written policy into dollars caught—maps directly onto policy intelligence as a product line rather than a services margin trap. Industrialization could shorten time-to-production for new edits after CMS releases, improve prepay attach rates for existing clients, and differentiate against pure analytics vendors that diagnose leakage without operationalizing fixes. Integration with Edifecs or similar stacks positions Cotiviti as the intelligence layer upstream of execution. Impact simulation against synthetic claims and CMS Physician Fee Schedule pricing (as demonstrated in the proof of concept) quantifies prepay ROI before deployment, strengthening client conversations.

Threats. Competitors and large payers may build internal LLM pipelines, reducing outsourcing appetite. Model errors that slip past weak SME controls could deploy incorrect edits at scale—a reputational and compliance catastrophe. CMS and state regulators may scrutinize AI-generated coding policy interpretations. Edifecs or cloud hyperscalers could move upstream into policy ingestion, compressing margin. Finally, conflating payment-integrity automation with clinical decision support invites regulatory pushback; the boundary must remain explicit in product design and marketing.""",
    ),
    (
        "Strategic Options",
        """Three concrete strategic paths emerge from the analysis and the proof-of-concept demonstration.

Option 1: Deep Edifecs agentic integration. Invest in a production-grade export and push adapter from the Policy Intelligence draft format into Edifecs' agentic deployment stack, preserving human-oversight checkpoints Edifecs emphasizes. Cotiviti becomes the authoritative policy-ingestion and change-detection front end; Edifecs remains execution. This leverages partner ecosystems, reduces duplicate rules-engine investment, and accelerates enterprise sales cycles where Edifecs is already licensed.

Option 2: Standalone Policy Intelligence product line. Productize the four-stage pipeline as a SaaS module for payers and revenue integrity teams: ingest CMS and payer PDFs, SME review queue, impact simulation, audit trail. Price on documents processed, edits approved, or prepay dollars modeled. Keeps optionality across rules-engine vendors via open JSON/YAML export. Requires investment in authentication, role-based SME enforcement, live CMS fetch, and structured NCCI table import—gaps noted in the Section 3 compliance review (~65% maturity in demo scope).

Option 3: Phase 2 VBP contract extension. Extend agents and rule patterns to payer–provider contracts: quality gates, shared-savings tiers, reconciliation stubs. Positions Cotiviti in contract administration—a growth vector as fee-for-service editing matures. Sequencing matters: establish billing-policy credibility first, then cross-sell VBP settlement logic to clients already using prepay edits.

Recommended posture: pursue Option 1 and Option 2 in parallel—integration for enterprise accounts, standalone module for mid-market—with Option 3 as a twelve-to-eighteen-month extension once SME workflow and export harden.""",
    ),
    (
        "Conclusion",
        """Policy intelligence industrializes Cotiviti's core competency for an era of AI-augmented, prepay-first payment integrity. The Policy Intelligence Pipeline proof of concept demonstrates feasibility across ingest, diff, draft, SME review, and impact simulation, while honestly scoping production gaps. Cotiviti's strategic choice is not whether to automate policy reading, but how to productize and govern that automation without surrendering the SME accountability that Medicare and commercial payers require.""",
    ),
]

REFERENCES = [
    "Centers for Medicare & Medicaid Services. (n.d.). National Correct Coding Initiative. U.S. Department of Health and Human Services. https://www.cms.gov/medicare/coding-billing/national-correct-coding-initiative-ncci-edits",
    "Centers for Medicare & Medicaid Services. (2025). Medicare NCCI policy manual (2025 ed.). https://www.cms.gov/files/document/2025nccimedicarepolicymanualcompletepdf.pdf",
    "Cotiviti. (2026). Policy Intelligence Pipeline (Version hackathon POC) [Computer software]. https://github.com/cotiviti/policy-intelligence-pipeline",
    "Edifecs. (2025). Agentic deployment with human oversight [Product brief]. https://www.edifecs.com",
    "HFMA. (2024). Payment integrity trends: Prepayment vs. post-payment recovery. Healthcare Financial Management Association. https://www.hfma.org",
    "McKinsey & Company. (2024). Value-based care: Scaling what works in healthcare payment reform. https://www.mckinsey.com/industries/healthcare",
    "NAHRI. (2025, June). CMS publishes 2025 Q3 update for NCCI edit files. NAHRI. https://nahri.org/articles/cms-publishes-2025-q3-update-ncci-edit-files",
    "Rajkomar, A., Dean, J., & Kohane, I. (2019). Machine learning in medicine. New England Journal of Medicine, 380(14), 1347–1358. https://doi.org/10.1056/NEJMra1814259",
]


def set_apa_paragraph(paragraph, text: str, *, bold: bool = False, size: int = 12) -> None:
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold


def main() -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    title = doc.add_paragraph()
    title.alignment = 1  # center
    set_apa_paragraph(title, TITLE, bold=True)

    meta = doc.add_paragraph()
    meta.alignment = 1
    set_apa_paragraph(meta, "Cotiviti Critical Assessment — Written Report\nJune 30, 2026", size=12)

    doc.add_paragraph()

    for heading, body in SECTIONS:
        h = doc.add_paragraph()
        set_apa_paragraph(h, heading, bold=True)
        for para in body.strip().split("\n\n"):
            p = doc.add_paragraph()
            set_apa_paragraph(p, para.strip())

    doc.add_page_break()
    ref_title = doc.add_paragraph()
    set_apa_paragraph(ref_title, "References", bold=True)
    for ref in REFERENCES:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        set_apa_paragraph(p, ref)

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()

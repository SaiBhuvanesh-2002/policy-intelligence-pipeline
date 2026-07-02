"""Extract plain text from uploaded files (PDF / DOCX / TXT / CSV).

Reuses the PDF extractor from the fetcher so there is a single code path for
turning bytes into policy text, regardless of how the bytes arrived (upload,
URL fetch, or seed).
"""
from __future__ import annotations

import io

from .fetcher import _extract_pdf


class UnsupportedFileType(Exception):
    pass


def _extract_docx(content: bytes) -> str:
    import docx  # python-docx

    document = docx.Document(io.BytesIO(content))
    parts: list[str] = [p.text for p in document.paragraphs if p.text and p.text.strip()]
    # Include simple table cell text (payer policies often use tables).
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_text(content: bytes) -> str:
    return content.decode("utf-8", errors="replace")


def extract_text(filename: str, content: bytes) -> str:
    """Dispatch on file extension / signature and return extracted text."""
    name = (filename or "").lower()
    if name.endswith(".pdf") or content[:5] == b"%PDF-":
        return _extract_pdf(content)
    if name.endswith(".docx"):
        return _extract_docx(content)
    if name.endswith((".txt", ".csv", ".md", ".text")):
        return _extract_text(content)
    # Fallback: try utf-8 decode for unknown but likely-text payloads.
    try:
        return _extract_text(content)
    except Exception as exc:  # pragma: no cover - defensive
        raise UnsupportedFileType(
            f"Unsupported file type for '{filename}'. Use PDF, DOCX, TXT, or CSV."
        ) from exc
